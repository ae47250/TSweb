import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def latest_synthetic_json(root):
    candidates = sorted(
        root.glob("synthetic-messy-input-review-data-*.json"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise SystemExit("No synthetic messy input review JSON files found.")
    return candidates[0]


def set_table_borders(table, color="DADCE0", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_label_run(paragraph, label):
    run = paragraph.add_run(label)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(67, 67, 67)
    return run


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, before, after in [
        ("Heading 1", 16, 14, 6),
        ("Heading 2", 12, 10, 4),
    ]:
      style = doc.styles[name]
      style.font.name = "Arial"
      style.font.size = Pt(size)
      style.font.color.rgb = RGBColor(0, 0, 0)
      style.paragraph_format.space_before = Pt(before)
      style.paragraph_format.space_after = Pt(after)


def add_case_table(doc, record):
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    set_table_borders(table)

    widths = [1.15, 5.35]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    labels = ["Case", "Input text", "TD2 job summary"]
    values = [
        f"{record['id']} ({'PASS' if record.get('pass') else 'FAIL'})",
        record.get("input", ""),
        record.get("actual", {}).get("td2_job_summary", ""),
    ]

    for index, (label, value) in enumerate(zip(labels, values)):
        label_cell, value_cell = table.rows[index].cells
        shade_cell(label_cell, "F2F4F7")
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        add_label_run(label_p, label)

        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        run = value_p.add_run(str(value))
        run.font.size = Pt(9)
        if label == "TD2 job summary":
            run.bold = True

    doc.add_paragraph()


def category_from_id(case_id):
    return re.sub(r"-\d{3}-(?:base|extra-spaces|followup-label|speech|td-short)$", "", case_id)


def main():
    root = Path.cwd()
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else latest_synthetic_json(root)
    payload = json.loads(source.read_text(encoding="utf-8"))
    stamp = payload.get("timestamp") or source.stem.replace("synthetic-messy-input-review-data-", "")
    out = root / "reports" / f"Synthetic TD2 Input and Job Summary {stamp}.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Synthetic Messy Input TD2 Job Summary Review")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(20)
    title_run.bold = True

    meta = doc.add_paragraph()
    meta.add_run(
        f"Source: {source.name}. Total: {payload['summary']['total']}; "
        f"Passed: {payload['summary']['passed']}; Failed: {payload['summary']['failed']}."
    )

    note = doc.add_paragraph()
    note.add_run("TD2 job summary lines are bold.").italic = True

    categories = {}
    for record in payload["records"]:
        category = category_from_id(record["id"])
        categories.setdefault(category, []).append(record)

    for category, records in categories.items():
        heading = doc.add_paragraph(style="Heading 1")
        heading.add_run(f"{category} ({len(records)})")
        for record in records:
            add_case_table(doc, record)

    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
