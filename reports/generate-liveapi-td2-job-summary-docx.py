from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports" / "LIVEapi-results-2026-07-01_19-55-50cases.jsonl"
OUT = ROOT / "reports" / "LIVEapi-TD2-job-summary-review-numbered-2026-07-01_19-55-50cases.docx"
PREVIEW = ROOT / "reports" / "LIVEapi-TD2-job-summary-review-numbered-preview.png"
CARD_OUT = ROOT / "reports" / "LIVEapi-TD2-job-summary-cards-2026-07-01_19-55-50cases.docx"
CARD_PREVIEW = ROOT / "reports" / "LIVEapi-TD2-job-summary-cards-preview.png"
GUIDE_OUT = ROOT / "reports" / "LIVEapi-TD2-job-summary-what-to-look-for.docx"
GUIDE_PREVIEW = ROOT / "reports" / "LIVEapi-TD2-job-summary-what-to-look-for-preview.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="000000", size="18") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.twips)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_text(paragraph, text: str, *, bold=False, color="111827", size=8.0) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text or "")
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def load_rows():
    rows = []
    with SOURCE.open("r", encoding="utf-8") as handle:
        for line in handle:
            if not line.strip():
                continue
            item = json.loads(line)
            case_number = len(rows) + 1
            rows.append(
                {
                    "case_number": case_number,
                    "input": item.get("raw_input", ""),
                    "summary": (item.get("td2_rendered_fields") or {}).get("job_summary", ""),
                }
            )
    return rows


def get_font(size: int, bold: bool = False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_for_width(draw: ImageDraw.ImageDraw, text: str, font, width: int, max_lines: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
        if len(lines) >= max_lines:
            break
    if current and len(lines) < max_lines:
        lines.append(current)
    if len(lines) == max_lines and len(" ".join(lines)) < len(text):
        lines[-1] = lines[-1].rstrip(".,; ") + "..."
    return lines or [""]


def draw_wrapped(draw: ImageDraw.ImageDraw, xy, text: str, font, fill, width: int, line_gap: int, max_lines: int):
    x, y = xy
    for line in wrap_for_width(draw, text, font, width, max_lines):
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap


def create_preview(rows) -> None:
    image = Image.new("RGB", (1800, 1260), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(26, bold=True)
    subtitle_font = get_font(15)
    header_font = get_font(17, bold=True)
    body_font = get_font(15)

    margin = 42
    input_x = margin
    summary_x = 1070
    input_w = 990
    summary_w = 680
    row_h = 132
    top = 105

    draw.text((margin, 24), "LIVEapi TD2 Job Summary Review", font=title_font, fill="#111827")
    draw.text(
        (margin, 60),
        "Preview of the Word document layout: raw input text on the left, TD2 job summary on the right.",
        font=subtitle_font,
        fill="#4B5563",
    )

    draw.rounded_rectangle((input_x, top, summary_x + summary_w, top + 42), radius=4, fill="#EDE9FE", outline="#C4B5FD")
    draw.line((summary_x - 16, top, summary_x - 16, top + 42 + row_h * 8), fill="#D1D5DB", width=2)
    draw.text((input_x + 14, top + 11), "Input text", font=header_font, fill="#3B0764")
    draw.text((summary_x, top + 11), "TD2 job summary", font=header_font, fill="#3B0764")

    y = top + 42
    for index, row in enumerate(rows[:8], start=1):
        fill = "#FFFFFF" if index % 2 else "#F9FAFB"
        draw.rectangle((input_x, y, summary_x + summary_w, y + row_h), fill=fill, outline="#E5E7EB")
        draw.line((summary_x - 16, y, summary_x - 16, y + row_h), fill="#E5E7EB", width=2)
        draw.text((input_x + 14, y + 10), f"Case {row['case_number']}", font=header_font, fill="#111827")
        draw.text((summary_x, y + 10), f"Case {row['case_number']}", font=header_font, fill="#111827")
        draw_wrapped(draw, (input_x + 14, y + 38), row["input"], body_font, "#111827", input_w - 30, 4, 5)
        draw_wrapped(draw, (summary_x, y + 38), row["summary"], body_font, "#111827", summary_w - 12, 4, 5)
        y += row_h

    draw.text(
        (margin, 1218),
        "Preview shows first 8 of 50 rows. Full Word file contains all 50 cases.",
        font=subtitle_font,
        fill="#6B7280",
    )
    image.save(PREVIEW)


def add_run_line(paragraph, text: str, *, bold=False, color="111827", size=7.5) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_card(cell, row_data) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_width(cell, 5.05)
    set_cell_margins(cell, 120, 140, 120, 140)
    set_cell_border(cell, "000000", "18")
    set_cell_shading(cell, "FFFFFF")

    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(3)
    add_run_line(p, f"Case {row_data['case_number']}", bold=True, size=8.5)

    input_label = cell.add_paragraph()
    input_label.paragraph_format.space_before = Pt(1)
    input_label.paragraph_format.space_after = Pt(1)
    add_run_line(input_label, "Input text", bold=True, color="3B0764", size=7.5)

    input_text = cell.add_paragraph()
    input_text.paragraph_format.space_after = Pt(5)
    input_text.paragraph_format.line_spacing = 1.05
    add_run_line(input_text, row_data["input"], size=7.0)

    summary_label = cell.add_paragraph()
    summary_label.paragraph_format.space_before = Pt(2)
    summary_label.paragraph_format.space_after = Pt(1)
    add_run_line(summary_label, "TD2 job summary", bold=True, color="3B0764", size=7.5)

    summary_text = cell.add_paragraph()
    summary_text.paragraph_format.line_spacing = 1.05
    add_run_line(summary_text, row_data["summary"], size=7.0)


def create_card_doc(rows) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(7)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_paragraph_text(title, "LIVEapi TD2 Job Summary Cards", bold=True, size=11)

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for index in range(0, len(rows), 2):
        doc_row = table.add_row()
        doc_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        add_card(doc_row.cells[0], rows[index])
        if index + 1 < len(rows):
            add_card(doc_row.cells[1], rows[index + 1])
        else:
            set_cell_width(doc_row.cells[1], 5.05)
            set_cell_margins(doc_row.cells[1], 120, 140, 120, 140)

    doc.save(CARD_OUT)


def create_card_preview(rows) -> None:
    image = Image.new("RGB", (1800, 1260), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(26, bold=True)
    label_font = get_font(16, bold=True)
    body_font = get_font(14)
    case_font = get_font(18, bold=True)

    draw.text((42, 24), "LIVEapi TD2 Job Summary Cards", font=title_font, fill="#111827")
    card_w = 835
    card_h = 255
    gap = 32
    left_x = 42
    right_x = left_x + card_w + gap
    y0 = 78

    for idx, row in enumerate(rows[:8]):
        x = left_x if idx % 2 == 0 else right_x
        y = y0 + (idx // 2) * (card_h + 24)
        draw.rectangle((x, y, x + card_w, y + card_h), fill="#FFFFFF", outline="#000000", width=5)
        draw.text((x + 18, y + 16), f"Case {row['case_number']}", font=case_font, fill="#111827")
        draw.text((x + 18, y + 48), "Input text", font=label_font, fill="#3B0764")
        draw_wrapped(draw, (x + 18, y + 74), row["input"], body_font, "#111827", card_w - 36, 3, 4)
        draw.text((x + 18, y + 151), "TD2 job summary", font=label_font, fill="#3B0764")
        draw_wrapped(draw, (x + 18, y + 177), row["summary"], body_font, "#111827", card_w - 36, 3, 4)

    draw.text((42, 1218), "Preview shows first 8 of 50 numbered cards. Full Word file contains all 50 cases.", font=get_font(15), fill="#6B7280")
    image.save(CARD_PREVIEW)


def guide_items() -> list[str]:
    return [
        "Completeness - does the summary capture the actual tree work Tree Dude typed?",
        "Tree count - does the number of trees match the input, especially messy counts?",
        "Species - does it preserve tree type and mixed species instead of flattening them?",
        "Work action - does it say remove, trim, cut, haul, or cleanup correctly?",
        "Location detail - does it keep useful job location details like garage, roof, shed, or driveway?",
        "Coherence - does it read like plain English instead of fragments or repeated text?",
        "No leakage - does it avoid phone, email, service address, prices, and option labels in the job summary?",
        "No invention - does it avoid adding facts that were not in the raw input?",
        "Missing-info behavior - if the input is vague, does TD2 make the uncertainty visible?",
        "Customer readiness - would the summary be acceptable after Tree Dude reviews and confirms it?",
    ]


def create_guide_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    set_paragraph_text(title, "What To Look For In TD2 Job Summaries", bold=True, size=16)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    set_paragraph_text(subtitle, "Use these one-line hints while comparing each raw input to the TD2 job summary.", color="4B5563", size=10)

    for item in guide_items():
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(5)
        add_run_line(p, "- ", bold=True, size=11)
        category, hint = item.split(" - ", 1)
        add_run_line(p, category, bold=True, size=11)
        add_run_line(p, f" - {hint}", size=11)

    doc.save(GUIDE_OUT)


def create_guide_preview() -> None:
    image = Image.new("RGB", (1200, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = get_font(30, bold=True)
    body_font = get_font(22)
    bold_font = get_font(22, bold=True)

    x = 54
    y = 40
    draw.text((x, y), "What To Look For In TD2 Job Summaries", font=title_font, fill="#111827")
    y += 55
    draw.text((x, y), "One-line hints for comparing raw input to TD2 job summary.", font=get_font(18), fill="#4B5563")
    y += 52

    for item in guide_items():
        category, hint = item.split(" - ", 1)
        draw.text((x, y), "- ", font=bold_font, fill="#111827")
        draw.text((x + 26, y), category, font=bold_font, fill="#111827")
        category_width = draw.textbbox((0, 0), category, font=bold_font)[2]
        draw.text((x + 34 + category_width, y), f" - {hint}", font=body_font, fill="#111827")
        y += 58
    image.save(GUIDE_PREVIEW)


def main() -> None:
    rows = load_rows()

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    set_paragraph_text(title, "LIVEapi TD2 Job Summary Review", bold=True, color="111827", size=11)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(5)
    set_paragraph_text(
        subtitle,
        "50 live API smoke-test cases. Only raw input text and the TD2 job summary are shown.",
        color="4B5563",
        size=7.5,
    )

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"

    input_width = 6.2
    summary_width = 4.25
    header_cells = table.rows[0].cells
    for cell, text, width in (
        (header_cells[0], "Input text", input_width),
        (header_cells[1], "TD2 job summary", summary_width),
    ):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, width)
        set_cell_margins(cell, 80, 100, 80, 100)
        set_cell_shading(cell, "EDE9FE")
        p = cell.paragraphs[0]
        p.text = ""
        set_paragraph_text(p, text, bold=True, color="3B0764", size=8.5)
    set_repeat_table_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for cell, text, width in (
            (row.cells[0], f"Case {row_data['case_number']}\n{row_data['input']}", input_width),
            (row.cells[1], f"Case {row_data['case_number']}\n{row_data['summary']}", summary_width),
        ):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, width)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.text = ""
            set_paragraph_text(p, text, size=7.5)

    doc.save(OUT)
    create_preview(rows)
    create_card_doc(rows)
    create_card_preview(rows)
    create_guide_doc()
    create_guide_preview()
    print(OUT)
    print(PREVIEW)
    print(CARD_OUT)
    print(CARD_PREVIEW)
    print(GUIDE_OUT)
    print(GUIDE_PREVIEW)


if __name__ == "__main__":
    main()
